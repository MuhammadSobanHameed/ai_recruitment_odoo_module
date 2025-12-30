# ai_extractor.py
import pytesseract
from PIL import Image, ImageEnhance
import cv2
import numpy as np
import json
import time
import io
import re
import logging
import tempfile
import os
from typing import Dict, Any, List
import groq


# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Import required dependencies
try:
    import fitz  # PyMuPDF
    import groq
    from docx import Document
    from dateutil import parser as date_parser  # For better date parsing
except ImportError as e:
    logger.error(f"Missing dependency: {e}")
    logger.info("Install required packages: pip install pymupdf python-docx groq python-dateutil")


class AIExtractor:
    """
    Extract structured information from CV/Resume using Groq AI API
    """
    
    def __init__(self, api_key=None):
        """Initialize the extractor with API key"""
        if api_key:
            api_key = api_key.strip()
        self.client = groq.Client(api_key=api_key)
        
    def extract_from_binary(self, binary_data):
        """
        Main entry point - extract text from binary file data and process with AI
        
        Args:
            binary_data: Binary content of uploaded CV file
            
        Returns:
            dict: Structured CV data in JSON format
        """
        try:
            # Extract text from binary data
            text = self._extract_text_from_binary(binary_data)
            if not text or len(text.strip()) < 50:  # Check for meaningful content
                logger.error("Insufficient text content extracted from document")
                return {"error": "Could not extract sufficient text from document"}

            # Process the extracted text with the AI
            extracted_data = self.extract_info_with_llama(text)
            if not extracted_data:
                logger.error("AI extraction returned empty result")
                return {"error": "AI processing failed to extract data"}
                
            # Post-process the data to ensure it matches expected structure
            cleaned_data = self._clean_extracted_data(extracted_data)
            
            # Extract detailed suggested summary if not already present
            if not cleaned_data.get('Personal_Info', {}).get('Suggested_Summary') or len(cleaned_data.get('Personal_Info', {}).get('Suggested_Summary', '')) < 100:
                cleaned_data = self._enhance_suggested_summary(text, cleaned_data)

            # Infer gender from context if not present
            if not cleaned_data.get("Personal_Info", {}).get("Gender"):
                name = cleaned_data.get("Personal_Info", {}).get("Full Name", "")
                summary = cleaned_data.get("Personal_Info", {}).get("Suggested_Summary", "")
                inferred_gender = self._infer_gender_from_context(name, summary)
                if inferred_gender:
                    if "Personal_Info" not in cleaned_data:
                        cleaned_data["Personal_Info"] = {}
                    cleaned_data["Personal_Info"]["Gender"] = inferred_gender

            logger.info(f"Successfully extracted {len(cleaned_data)} data points")

            return cleaned_data


        except Exception as e:
            logger.error(f"Error in extraction process: {str(e)}", exc_info=True)
            return {"error": f"Extraction failed: {str(e)}"}

    def _infer_gender_from_context(self, name: str, summary: str) -> str:
        """Use LLM to guess gender based on name and summary context"""
        try:
            prompt = f"""Determine the most likely gender of the person based on their name and context.

    Name: {name}
    Summary: {summary}

    Respond with only one of the following: "male", "female", or "other".
    """
            response = self.client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2,
                max_tokens=5
            )
            gender = response.choices[0].message.content.strip().lower()
            if gender in ["male", "female", "other"]:
                return gender
        except Exception as e:
            logger.warning(f"Failed to infer gender from context: {str(e)}")
        return None

    def _extract_text_from_binary(self, binary_data):
        """
        Extract text from various document formats, including CamScanner PDFs
        
        Args:
            binary_data: Binary content of uploaded CV file
                
        Returns:
            str: Extracted text content
        """
        try:
            # First try PDF (most common format)
            if binary_data[:4] == b'%PDF' or b'%PDF-' in binary_data[:1024]:
                try:
                    with fitz.open(stream=binary_data, filetype="pdf") as doc:
                        text = ""
                        # Enhanced PDF extraction with text blocks for better layout handling
                        for page_num in range(len(doc)):
                            page = doc[page_num]
                            # Extract text blocks to maintain structure better
                            blocks = page.get_text("blocks")
                            for block in blocks:
                                text += block[4] + "\n"
                            
                            # If blocks method doesn't return enough text, fallback to standard
                            if len(text.strip()) < 100:
                                text += page.get_text() + "\n"
                        
                        # Enhanced table detection for CVs - look for common patterns
                        if "CURRICULUM VITAE" in text or "CV" in text or "Resume" in text:
                            # Try to get a more detailed extraction for structured CV documents
                            detailed_text = ""
                            for page_num in range(len(doc)):
                                page = doc[page_num]
                                # Dictionary mode can sometimes preserve structure better
                                detailed_text += page.get_text("dict") + "\n"
                                # Also try raw mode to catch any missed content
                                detailed_text += page.get_text("rawdict") + "\n"
                            
                            # Combine both methods if detailed extraction gave meaningful results
                            if len(detailed_text.strip()) > 100:
                                text = text + "\n\n" + detailed_text
                        
                        # Check if this might be a CamScanner PDF or image-based PDF
                        if len(text.strip()) < 100 or "CamScanner" in text:
                            logger.info("Detected potential CamScanner or image-based PDF, attempting OCR")
                            try:
                                # Try to import OCR-related libraries

                                
                                ocr_text = ""
                                for page_num in range(len(doc)):
                                    page = doc[page_num]
                                    # Higher resolution for better OCR
                                    pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
                                    img_bytes = pix.tobytes("png")
                                    
                                    # Process with OCR
                                    img = Image.open(io.BytesIO(img_bytes))
                                    img = img.convert('L')  # Convert to grayscale
                                    
                                    # Pre-process image to improve OCR accuracy
                                    try:
                                        # Increase contrast
                                        enhancer = ImageEnhance.Contrast(img)
                                        img = enhancer.enhance(1.5)
                                        
                                        # Apply adaptive thresholding
                                        img_np = np.array(img)
                                        img_np = cv2.adaptiveThreshold(
                                            img_np, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                            cv2.THRESH_BINARY, 11, 2
                                        )
                                        img = Image.fromarray(img_np)
                                    except Exception as img_err:
                                        logger.warning(f"Image preprocessing failed: {img_err}, continuing with basic OCR")
                                    
                                    # OCR with improved settings
                                    try:
                                        custom_config = r'--oem 3 --psm 6'
                                        page_text = pytesseract.image_to_string(img, lang='eng', config=custom_config)
                                        ocr_text += page_text + "\n\n"
                                    except Exception as ocr_err:
                                        logger.warning(f"Detailed OCR failed: {ocr_err}, trying simpler approach")
                                        # Fallback to basic OCR
                                        try:
                                            page_text = pytesseract.image_to_string(img)
                                            ocr_text += page_text + "\n\n"
                                        except:
                                            pass
                                
                                if ocr_text.strip():
                                    logger.info(f"Successfully extracted {len(ocr_text)} characters with OCR")
                                    text = ocr_text  # Replace the minimal text with OCR results
                            except ImportError as imp_err:
                                logger.warning(f"OCR libraries not available: {imp_err}. Install with: pip install pytesseract pillow opencv-python")
                            except Exception as ocr_err:
                                logger.warning(f"OCR extraction failed: {ocr_err}")
                        
                        logger.info(f"Successfully extracted PDF text: {len(text)} characters")
                        return text.strip()
                except Exception as pdf_err:
                    logger.warning(f"PDF extraction failed: {pdf_err}")
                    # Try alternative PDF extraction method
                    try:
                        with fitz.open(stream=binary_data, filetype="pdf") as doc:
                            text = ""
                            for page in doc:
                                text += page.get_text("text") + "\n"
                            return text.strip()
                    except:
                        pass

                # Last resort for PDF: Try to extract embedded images and perform OCR
                try:
                    with fitz.open(stream=binary_data, filetype="pdf") as doc:
                        try:
                            import pytesseract
                            from PIL import Image
                            
                            text = ""
                            for page_num in range(len(doc)):
                                page = doc[page_num]
                                image_list = page.get_images(full=True)
                                
                                if image_list:
                                    for img_index, img_info in enumerate(image_list):
                                        xref = img_info[0]
                                        base_image = doc.extract_image(xref)
                                        image_bytes = base_image["image"]
                                        
                                        # Process the image with OCR
                                        img = Image.open(io.BytesIO(image_bytes))
                                        img_text = pytesseract.image_to_string(img)
                                        if img_text.strip():
                                            text += img_text + "\n"
                            
                            if text.strip():
                                logger.info(f"Extracted {len(text)} characters from embedded PDF images")
                                return text.strip()
                        except ImportError:
                            logger.warning("OCR libraries not available for image extraction")
                        except Exception as img_err:
                            logger.warning(f"Image extraction failed: {img_err}")
                except Exception as last_err:
                    logger.warning(f"Last resort PDF extraction failed: {last_err}")

            # Try DOCX format - improved to extract more content
            if binary_data[:4] == b'PK\x03\x04':  # ZIP signature (DOCX is ZIP-based)
                try:
                    doc = Document(io.BytesIO(binary_data))
                    
                    # Extract paragraph text
                    paragraphs = []
                    for p in doc.paragraphs:
                        if p.text.strip():
                            paragraphs.append(p.text.strip())
                    
                    # Improved table extraction for DOCX files - critical for CV formats
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                # Get text from paragraphs within cells (handles nested content)
                                cell_text = []
                                for paragraph in cell.paragraphs:
                                    if paragraph.text.strip():
                                        cell_text.append(paragraph.text.strip())
                                
                                # Also check for nested tables within cells
                                for nested_table in cell.tables:
                                    for nested_row in nested_table.rows:
                                        nested_text = []
                                        for nested_cell in nested_row.cells:
                                            if nested_cell.text.strip():
                                                nested_text.append(nested_cell.text.strip())
                                        if nested_text:
                                            cell_text.append(" | ".join(nested_text))
                                
                                # Add combined cell content
                                if cell_text:
                                    cell_content = " ".join(cell_text)
                                    # Clean up formatting markers that might appear in exported Word docs
                                    cell_content = re.sub(r'\{[^}]*\}', '', cell_content)
                                    cell_content = re.sub(r'\[(.*?)\](?:\{[^}]*\})?', r'\1', cell_content)
                                    row_text.append(cell_content)
                            
                            if row_text:
                                # Join with delimiter that preserves table structure
                                row_content = " | ".join(row_text)
                                # Clean common formatting artifacts from tables
                                row_content = re.sub(r'\{[^}]*\}', '', row_content)
                                row_content = re.sub(r'\[.*?\]', '', row_content)
                                row_content = re.sub(r'\.smallcaps', '', row_content)
                                paragraphs.append(row_content)
                    
                    # Extract headers/footers if possible
                    try:
                        for section in doc.sections:
                            # Try to get headers
                            for header in [section.header, section.first_page_header, section.even_page_header]:
                                if header and header.paragraphs:
                                    for p in header.paragraphs:
                                        if p.text.strip():
                                            paragraphs.append(p.text.strip())
                            
                            # Try to get footers
                            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                                if footer and footer.paragraphs:
                                    for p in footer.paragraphs:
                                        if p.text.strip():
                                            paragraphs.append(p.text.strip())
                    except:
                        # Header/footer extraction is not critical
                        pass
                    
                    # Join all paragraphs, then clean common formatting artifacts
                    text = '\n'.join(paragraphs)
                    # Clean LaTeX-like or Word export artifacts common in CV templates
                    text = re.sub(r'\\[a-z]+\{([^}]*)\}', r'\1', text)  # Clean LaTeX-style formatting
                    text = re.sub(r'\{\\[a-z]+\s', ' ', text)  # Clean RTF-style formatting starts
                    text = re.sub(r'\{[^}]*\}', '', text)  # Remove remaining curly braces content
                    text = re.sub(r'\[.*?\](?:\{[^}]*\})?', '', text)  # Remove square brackets content
                    text = re.sub(r'\.smallcaps', '', text)  # Remove .smallcaps markers
                    text = re.sub(r'\+\-+\+\-+\+\-+\+', '', text)  # Remove table formatting lines
                    text = re.sub(r'\={3,}', '', text)  # Remove table separator lines
                    
                    logger.info(f"Successfully extracted DOCX text: {len(text)} characters")
                    return text.strip()
                except Exception as docx_err:
                    logger.warning(f"DOCX extraction failed: {docx_err}")
            
            # Try RTF format
            if binary_data[:5] in [b'{\\rtf', b'{\\RTF']:
                # Save to temp file for processing
                with tempfile.NamedTemporaryFile(suffix='.rtf', delete=False) as temp:
                    temp.write(binary_data)
                    temp_path = temp.name
                
                try:
                    # Use striprtf if available
                    try:
                        from striprtf.striprtf import rtf_to_text
                        with open(temp_path, 'r', encoding='utf-8', errors='ignore') as rtf_file:
                            rtf_content = rtf_file.read()
                            text = rtf_to_text(rtf_content)
                            logger.info(f"Successfully extracted RTF text: {len(text)} characters")
                            return text.strip()
                    except ImportError:
                        logger.warning("striprtf not available, trying fallback method")
                        
                    # Manual RTF text extraction (basic)
                    with open(temp_path, 'r', encoding='utf-8', errors='ignore') as rtf_file:
                        content = rtf_file.read()
                        # Simple RTF parsing - remove RTF control codes
                        text = re.sub(r'\\[a-z0-9]+\s?', ' ', content)
                        text = re.sub(r'\{|\}|\\\n|\\\r|\\\t|\\\'[0-9a-f]{2}', '', text)
                        logger.info(f"Basic RTF extraction: {len(text)} characters")
                        return text.strip()
                        
                finally:
                    # Clean up temp file
                    if os.path.exists(temp_path):
                        os.unlink(temp_path)
            
            # Check if this could be an image file that contains CV text
            try:
                img_extensions = [b'JFIF', b'PNG', b'GIF', b'BMP', b'JPEG']
                is_image = False
                for sig in img_extensions:
                    if sig in binary_data[:20]:
                        is_image = True
                        break
                
                if is_image:
                    logger.info("Detected image file, attempting OCR")
                    try:
                        import pytesseract
                        from PIL import Image, ImageEnhance
                        
                        # Load the image
                        img = Image.open(io.BytesIO(binary_data))
                        
                        # Optimize for OCR
                        if img.mode != 'RGB':
                            img = img.convert('RGB')
                        
                        # Enhance image for better OCR results
                        enhancer = ImageEnhance.Contrast(img)
                        img = enhancer.enhance(1.5)
                        
                        # Perform OCR
                        text = pytesseract.image_to_string(img)
                        if text.strip():
                            logger.info(f"OCR extracted {len(text)} characters from image")
                            return text.strip()
                    except ImportError:
                        logger.warning("OCR libraries not available for image processing")
                    except Exception as img_ocr_err:
                        logger.warning(f"Image OCR failed: {img_ocr_err}")
            except:
                pass
            
            # Last resort: try as plain text with various encodings
            encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'windows-1252']
            for encoding in encodings:
                try:
                    text = binary_data.decode(encoding, errors='ignore')
                    if len(text.strip()) > 100:  # Only accept if meaningful content
                        logger.info(f"Extracted text using {encoding} encoding")
                        return text.strip()
                except UnicodeDecodeError:
                    continue
            
            logger.warning("Could not determine file format for text extraction")
            return binary_data.decode('utf-8', errors='ignore')  # Last attempt with error ignore
                
        except Exception as e:
            logger.error(f"Text extraction error: {str(e)}", exc_info=True)
            return ""
    
    def extract_info_with_llama(self, text):
        """
        Process text with Groq LLM to extract structured information
        
        Args:
            text: CV text content
            
        Returns:
            dict: Structured CV data
        """
        # Split text into manageable chunks to stay within API limits
        chunks = self.chunk_text(text)
        extracted_data = {}
        
        system_prompt = """You are an AI expert in CV/resume information extraction. Extract and organize all relevant information from the provided CV/resume.
        Return only valid JSON in this exact format:
        {
          "Personal_Info": {
            "Full Name": "string",
            "Suggested_Summary": "string", 
            "Summary": "string",
            "Date of Birth": "string",
            "Gender": "string",
            "Nationality": "string",
            "Marital Status": "string"
          },
          "Contact_Info": {
            "Email Address": "string",
            "Mobile Phone Number": "string",
            "Work Phone Number": "string",
            "Home Phone Number": "string"
          },
          "Location_Info": {
            "Full Address": "string",
            "City": "string",
            "Region/State": "string",
            "Country": "string",
            "Postal/Zip Code": "string"
          },
          "Social_Media": {
            "LinkedIn URL": "string",
            "Twitter/X URL": "string",
            "Facebook URL": "string",
            "GitHub URL": "string",
            "Portfolio Website": "string"
          },
          "Education": [
            {
              "Degree": "string",
              "Institution": "string",
              "Dates": "string",
              "Additional Info": "string"
            }
          ],
          "Professional Experience": [
            {
              "Job Titles": "string",
              "Companies": "string",
              "Employment Dates": "string",
              "Key Responsibilities": ["string"]
            }
          ],
          "Skills": {
            "Technical": ["string"],
            "Soft": ["string"],
            "Languages": ["string"],
            "Frameworks": ["string"],
            "Databases": ["string"],
            "Tools": ["string"],
            "Interests": ["string"],
            "Hobbies": ["string"]
          }
        }

        For the Suggested_Summary field, provide a detailed 3-5 sentence paragraph that highlights the candidate's professional background, key skills, experience level, and career achievements. This should be comprehensive but concise (150-300 words).
        
        Return ONLY valid JSON without explanation text, markdown or code blocks.
        Use null or empty arrays for missing information. Be precise and accurate.
        For dates, use any format that you find in the CV, don't convert them.
        """

        try:
            for i, chunk in enumerate(chunks):
                user_prompt = f"""Extract all relevant CV/resume information from this document text.
Extract exact dates, titles, responsibilities, and detailed contact information if present.
Pay special attention to creating a comprehensive Suggested_Summary that highlights the candidate's key qualifications, experience, and value proposition.
Remember to format phone numbers consistently.

DOCUMENT TEXT:
{chunk}
"""

                response = self.client.chat.completions.create(
                    model="llama-3.3-70b-versatile",  # Using latest available model
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    temperature=0.1,  # Low temperature for factual extraction
                    max_tokens=4096,
                    top_p=0.7,
                    stream=False
                )

                result = response.choices[0].message.content
                
                # Clean the result to extract valid JSON
                clean_json = self._extract_json_from_text(result)
                
                if not clean_json:
                    logger.warning(f"Failed to extract valid JSON from chunk {i+1}")
                    continue
                    
                try:
                    chunk_data = json.loads(clean_json)
                    if not extracted_data:
                        extracted_data = chunk_data
                    else:
                        self._merge_extracted_data(extracted_data, chunk_data)
                        
                    logger.info(f"Successfully processed chunk {i+1}/{len(chunks)}")
                        
                except json.JSONDecodeError as je:
                    logger.error(f"JSON parsing error: {je} in:\n{clean_json[:100]}...")
                    continue

                # Rate limiting for multiple chunks
                if i < len(chunks) - 1:
                    time.sleep(1)

            return extracted_data

        except Exception as e:
            logger.error(f"AI extraction error: {str(e)}", exc_info=True)
            return {}

    def _extract_summary_text(self, full_text):
        """
        Extract the existing summary/profile section from the CV text
        """
        # Common section headers that might contain the summary
        summary_headers = [
            r'(?:Professional\s+)?Summary',
            r'Profile',
            r'Personal\s+Profile',
            r'Professional\s+Profile',
            r'Career\s+Profile',
            r'Career\s+Summary',
            r'Executive\s+Summary',
            r'Overview',
            r'About\s+Me'
        ]
        
        # Create regex pattern to find summary sections
        pattern = '|'.join(f"({header})" for header in summary_headers)
        sections = full_text.split('\n\n')
        
        for i, section in enumerate(sections):
            # Look for summary section headers
            if re.search(f"^(?:{pattern})\\s*:?\\s*$", section.strip(), re.IGNORECASE):
                # If found, get the next section which should contain the actual summary
                if i + 1 < len(sections):
                    return sections[i + 1].strip()
                    
            # Also check for inline summary headers
            match = re.match(f"(?:{pattern})\\s*:?\\s*(.+)", section.strip(), re.IGNORECASE | re.DOTALL)
            if match:
                return match.group(1).strip()
        
        # If no explicit summary section found, look for summary-like content at the start
        first_section = sections[0] if sections else ""
        if len(first_section.strip()) > 50 and len(first_section.strip()) < 1000:
            # Check if it looks like a summary (no common headers, contact info, etc.)
            if not re.search(r'(education|experience|skills|contact|phone|email|address)', first_section, re.IGNORECASE):
                return first_section.strip()
        
        return ""

    def _enhance_suggested_summary(self, full_text, data):
        """
        Extract the suggested summary from the CV text
        
        Args:
            full_text: The full text of the CV
            data: The existing extracted data
            
        Returns:
            dict: Updated data with extracted summary
        """
        try:
            # Extract the actual summary from the CV text
            extracted_summary = self._extract_summary_text(full_text)
            
            if extracted_summary:
                # Clean up the extracted summary
                summary = re.sub(r'^(Suggested Summary:|Summary:|Profile:|About:|Overview:)\s*', '', extracted_summary, flags=re.IGNORECASE)
                summary = summary.strip()
                
                # Update the data with the extracted summary
                if 'Personal_Info' not in data:
                    data['Personal_Info'] = {}
                
                data['Personal_Info']['Suggested_Summary'] = summary
                data['Personal_Info']['Summary'] = summary[:150] if len(summary) > 150 else summary
                
            return data
            
        except Exception as e:
            logger.error(f"Error extracting suggested summary: {str(e)}")
            return data

    def _extract_json_from_text(self, text):
        """Extract valid JSON from text that might contain markdown or explanations"""
        # First try to find content within code blocks
        if "```json" in text:
            match = re.search(r'```json\s*([\s\S]*?)\s*```', text)
            if match:
                return match.group(1).strip()
        elif "```" in text:
            match = re.search(r'```\s*([\s\S]*?)\s*```', text)
            if match:
                return match.group(1).strip()
                
        # Try to find JSON by brackets if no code block
        try:
            start_idx = text.find('{')
            end_idx = text.rfind('}') + 1
            if start_idx >= 0 and end_idx > start_idx:
                potential_json = text[start_idx:end_idx]
                # Validate by trying to parse it
                json.loads(potential_json)
                return potential_json
        except json.JSONDecodeError:
            pass
            
        # Final attempt - find the largest {...} block
        json_blocks = re.findall(r'\{(?:[^{}]|(?:\{(?:[^{}]|(?:\{[^{}]*\}))*\}))*\}', text)
        if json_blocks:
            # Get the largest block as it's likely the main JSON object
            largest_block = max(json_blocks, key=len)
            try:
                # Validate
                json.loads(largest_block)
                return largest_block
            except:
                pass
                
        logger.error("Could not extract valid JSON from API response")
        return ""

    def chunk_text(self, text, max_length=4000):
        """Split text into chunks suitable for API processing"""
        if len(text) <= max_length:
            return [text]
            
        chunks = []
        current_chunk = ""
        
        # Try to split by natural section breaks first
        sections = re.split(r'(?:\r?\n){2,}', text)
        
        for section in sections:
            if len(current_chunk) + len(section) + 2 <= max_length:
                current_chunk += section + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                
                # If the section itself is too long, split by lines
                if len(section) > max_length:
                    lines = section.split('\n')
                    current_chunk = ""
                    for line in lines:
                        if len(current_chunk) + len(line) + 1 <= max_length:
                            current_chunk += line + "\n"
                        else:
                            if current_chunk:
                                chunks.append(current_chunk.strip())
                            
                            # If a single line is too long, split by words
                            if len(line) > max_length:
                                words = line.split()
                                current_chunk = ""
                                for word in words:
                                    if len(current_chunk) + len(word) + 1 <= max_length:
                                        current_chunk += word + " "
                                    else:
                                        chunks.append(current_chunk.strip())
                                        current_chunk = word + " "
                            else:
                                current_chunk = line + "\n"
                else:
                    current_chunk = section + "\n\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
            
        logger.info(f"Split text into {len(chunks)} chunks for processing")
        return chunks

    def _merge_extracted_data(self, main_data, new_data):
        """Merge data from multiple chunks intelligently"""
        # Helper function to select the better value
        def select_best_value(current, new):
            if not current:
                return new
            if not new:
                return current
                
            # Prefer longer, more detailed values
            if len(new) > len(current) * 1.5:  # Significantly more info
                return new
            return current
            
        # Special handling for Suggested Summary
        if 'Personal_Info' in new_data and 'Suggested_Summary' in new_data['Personal_Info']:
            if 'Personal_Info' not in main_data:
                main_data['Personal_Info'] = {}
            
            # If we already have a summary, choose the longer/better one
            if 'Suggested_Summary' in main_data['Personal_Info']:
                current_summary = main_data['Personal_Info']['Suggested_Summary']
                new_summary = new_data['Personal_Info']['Suggested_Summary']
                
                # Choose the better summary based on length and content
                if new_summary and (not current_summary or len(new_summary) > len(current_summary) * 1.2):
                    main_data['Personal_Info']['Suggested_Summary'] = new_summary
            else:
                # If no existing summary, add the new one
                main_data['Personal_Info']['Suggested_Summary'] = new_data['Personal_Info']['Suggested_Summary']
            
        # Merge list type fields (Education, Experience, etc.)
        for key in ['Education', 'Professional Experience']:
            if key in new_data and isinstance(new_data[key], list) and new_data[key]:
                # Initialize if not present
                if key not in main_data:
                    main_data[key] = []
                
                for item in new_data[key]:
                    # Check if this item might be a duplicate
                    is_duplicate = False
                    for existing in main_data[key]:
                        # For education, check institution and degree
                        if key == 'Education' and 'Institution' in item and 'Institution' in existing:
                            if item['Institution'] and existing['Institution'] and \
                            self._similarity_score(item['Institution'], existing['Institution']) > 0.7:
                                # Merge the entries
                                for field in existing:
                                    existing[field] = select_best_value(existing[field], item.get(field, ''))
                                is_duplicate = True
                                break
                        
                        # For experience, check company and job title
                        elif key == 'Professional Experience' and 'Companies' in item and 'Companies' in existing:
                            if item['Companies'] and existing['Companies'] and \
                            self._similarity_score(item['Companies'], existing['Companies']) > 0.7:
                                # Merge the entries
                                for field in existing:
                                    if field == 'Key Responsibilities' and isinstance(existing[field], list) and isinstance(item.get(field, []), list):
                                        # Merge responsibilities without duplicates
                                        existing_resp = set(existing[field])
                                        for resp in item.get(field, []):
                                            existing_resp.add(resp)
                                        existing[field] = list(existing_resp)
                                    else:
                                        existing[field] = select_best_value(existing[field], item.get(field, ''))
                                is_duplicate = True
                                break
                    
                    if not is_duplicate:
                        main_data[key].append(item)
        
        # Merge dictionary fields (Personal_Info, Contact_Info, etc.)
        for group_key in ['Personal_Info', 'Contact_Info', 'Location_Info', 'Social_Media']:
            if group_key in new_data and isinstance(new_data[group_key], dict):
                if group_key not in main_data:
                    main_data[group_key] = {}
                    
                for field, value in new_data[group_key].items():
                    # Skip Suggested_Summary as it's already handled
                    if field == 'Suggested_Summary':
                        continue
                        
                    if value and value != "N/A" and value != "None" and value != "Unknown":
                        existing_val = main_data[group_key].get(field)
                        main_data[group_key][field] = select_best_value(existing_val, value)
        
        # Special handling for Skills
        if 'Skills' in new_data and isinstance(new_data['Skills'], dict):
            if 'Skills' not in main_data:
                main_data['Skills'] = {}
                
            for skill_type, skills in new_data['Skills'].items():
                if skills and isinstance(skills, list):
                    valid_skills = [s for s in skills if s and s != "N/A" and s != "None"]
                    if valid_skills:
                        # FIX: Initialize empty list if skill_type doesn't exist in main_data['Skills']
                        if skill_type not in main_data['Skills']:
                            main_data['Skills'][skill_type] = []
                        
                        existing = main_data['Skills'][skill_type]
                        
                        # FIX: Ensure existing is always a list
                        if existing is None:
                            existing = []
                            main_data['Skills'][skill_type] = existing
                            
                        for skill in valid_skills:
                            if skill not in existing:
                                existing.append(skill)
        
        return main_data
            
    def _similarity_score(self, str1, str2):
        """Calculate simple similarity between two strings"""
        # Convert to lowercase and remove common punctuation
        s1 = re.sub(r'[^\w\s]', '', str1.lower())
        s2 = re.sub(r'[^\w\s]', '', str2.lower())
        
        # Count matching words
        words1 = set(s1.split())
        words2 = set(s2.split())
        
        if not words1 or not words2:
            return 0
            
        common = words1.intersection(words2)
        return len(common) / max(len(words1), len(words2))
    
    def _clean_extracted_data(self, data):
        """
        Clean and normalize extracted data to ensure consistency
        """
        if not data:
            return {}
            
        # Format phone numbers consistently
        if 'Contact_Info' in data and isinstance(data['Contact_Info'], dict):
            for field in ['Mobile Phone Number', 'Work Phone Number', 'Home Phone Number']:
                if field in data['Contact_Info'] and data['Contact_Info'][field]:
                    # Keep only digits, plus sign, and parentheses
                    phone = data['Contact_Info'][field]
                    digits = re.sub(r'[^\d+()]', '', phone)
                    if digits:
                        data['Contact_Info'][field] = digits
        
        # Clean URLs
        if 'Social_Media' in data and isinstance(data['Social_Media'], dict):
            for field in ['LinkedIn URL', 'Twitter/X URL', 'Facebook URL', 'GitHub URL', 'Portfolio Website']:
                if field in data['Social_Media'] and data['Social_Media'][field]:
                    url = data['Social_Media'][field].strip()
                    # Add https:// if missing
                    if url and not url.startswith(('http://', 'https://')):
                        if not url.startswith('www.'):
                            url = 'www.' + url
                        url = 'https://' + url
                    data['Social_Media'][field] = url
        
        # Normalize and format dates for Odoo compatibility
        if 'Personal_Info' in data and 'Date of Birth' in data['Personal_Info']:
            dob = data['Personal_Info']['Date of Birth']
            if dob:
                try:
                    # Convert various date formats to Odoo's expected format '%Y-%m-%d'
                    parsed_date = None
                    
                    # Try DD-MM-YYYY format (like '03-01-1996')
                    if re.match(r'\d{1,2}-\d{1,2}-\d{4}', dob):
                        try:
                            parts = dob.split('-')
                            if len(parts) == 3:
                                day, month, year = parts
                                parsed_date = f"{year}-{month.zfill(2)}-{day.zfill(2)}"
                        except:
                            pass

                    # Try YYYY-MM-DD format (already correct)
                    if not parsed_date and re.match(r'\d{4}-\d{1,2}-\d{1,2}', dob):
                        try:
                            # Ensure proper padding of month and day
                            year, month, day = dob.split('-')
                            parsed_date = f"{year}-{month.zfill(2)}-{day.zfill(2)}"
                        except:
                            parsed_date = dob  # Keep as is if it's already in the right format
                    
                    # Try text date formats like "January 15, 1990" using dateutil
                    if not parsed_date:
                        try:
                            dt = date_parser.parse(dob, fuzzy=True)
                            parsed_date = dt.strftime("%Y-%m-%d")
                        except:
                            pass
                    
                    # Update only if we successfully parsed the date
                    if parsed_date:
                        data['Personal_Info']['Date of Birth'] = parsed_date
                    else:
                        # If we can't parse it, remove it to avoid errors
                        data['Personal_Info']['Date of Birth'] = None
                except Exception as e:
                    # If date parsing fails, set to None to avoid errors
                    data['Personal_Info']['Date of Birth'] = None
                    logger.warning(f"Failed to parse date: {dob}. Error: {str(e)}")
        
        # Handle and clean education data
        if 'Education' in data and isinstance(data['Education'], list):
            for edu in data['Education']:
                # Ensure all education entries have required fields
                for field in ['Degree', 'Institution', 'Dates', 'Additional Info']:
                    if field not in edu:
                        edu[field] = ""
        
        # Handle and clean experience data
        if 'Professional Experience' in data and isinstance(data['Professional Experience'], list):
            for exp in data['Professional Experience']:
                # Ensure all experience entries have required fields
                for field in ['Job Titles', 'Companies', 'Employment Dates']:
                    if field not in exp:
                        exp[field] = ""
                
                # Ensure 'Key Responsibilities' is a list
                if 'Key Responsibilities' not in exp:
                    exp['Key Responsibilities'] = []
                elif isinstance(exp['Key Responsibilities'], str):
                    # Convert string to list
                    exp['Key Responsibilities'] = [exp['Key Responsibilities']]
                
        return data